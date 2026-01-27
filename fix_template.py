from docx import Document
from docx.shared import Pt
from copy import deepcopy

# Open the original template
doc = Document(r'C:\Users\theul\Downloads\indigenous_responses_template.docx')

# Answers
answer_a = """Indigenous people adopted Western education through missionary schools and colonial educational institutions. Many among the colonial elite enrolled in these Western-style schools, where they learned European languages, Enlightenment values such as liberty, equality, and democracy, as well as Western academic subjects and religious teachings. This education created a new class of Western-educated indigenous elites who often saw themselves as modernizers capable of leading the regeneration of their societies. For example, in India, many upper-class Indians attended British schools and universities, absorbing ideas that they would later use to advocate for self-governance."""

answer_b = """Europeans used the concept of tribalism as a tool to facilitate colonial administration and control over African populations. Colonial governments required people to identify their "tribe" on applications for jobs, schools, and identity cards, which artificially created and reinforced tribal divisions that had not previously existed in the same rigid form. This served European interests in multiple ways: it kept African populations divided and easier to manage, reinforced European beliefs in African "primitiveness," and allowed colonizers to practice divide-and-rule strategies by favoring certain groups over others. The idea of an Africa sharply divided into separate and distinct "tribes" was largely a European invention that reflected their racial hierarchies and made the administration of vast colonial territories more manageable."""

answer_c = """AFRICA: The Maji Maji Rebellion (1905-1907) in German East Africa (present-day Tanzania) represented a violent uprising against European colonial rule. Indigenous peoples, frustrated by forced cotton cultivation and brutal colonial policies, united across ethnic lines to resist German authority. The rebels believed that sacred water ("maji" in Swahili) blessed by spiritual leaders would protect them from German bullets. Although the rebellion was eventually crushed with devastating consequences, it demonstrated the widespread resentment of European exploitation and the willingness of colonized peoples to resist, even at great cost.

ASIA: In India, the formation of the Indian National Congress in 1885 represented a more organized, political form of resistance to British imperialism. Western-educated Indians used the very Enlightenment ideals they had learned in colonial schools—such as self-determination, representative government, and national sovereignty—to challenge British rule. Initially advocating for greater Indian participation in governance, the Congress eventually grew into a powerful independence movement. This response demonstrated how colonized peoples could strategically adopt and adapt Western political concepts to advocate for their own rights and ultimately achieve independence."""

# Find and modify the Name field
for para in doc.paragraphs:
    if para.text.strip().startswith("Name:"):
        # Clear and set new text
        para.clear()
        para.add_run("Name: Andrew Weir")
        break

# Now we need to find the tables and put answers in them
tables = doc.tables
print(f"Found {len(tables)} tables")

# The template has 3 tables for the 3 response boxes
if len(tables) >= 3:
    # Table 0 - Answer a
    cell = tables[0].cell(0, 0)
    # Clear existing content
    for p in cell.paragraphs:
        p.clear()
    cell.paragraphs[0].add_run(answer_a)
    
    # Table 1 - Answer b  
    cell = tables[1].cell(0, 0)
    for p in cell.paragraphs:
        p.clear()
    cell.paragraphs[0].add_run(answer_b)
    
    # Table 2 - Answer c
    cell = tables[2].cell(0, 0)
    for p in cell.paragraphs:
        p.clear()
    cell.paragraphs[0].add_run(answer_c)

# Remove the old answers that were outside the boxes
# Find paragraphs containing answer text and clear them
paras_to_clear = []
for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    if text.startswith("Indigenous people adopted Western education"):
        paras_to_clear.append(i)
    elif text.startswith("Europeans used the concept of tribalism"):
        paras_to_clear.append(i)
    elif text.startswith("AFRICA: The Maji Maji"):
        paras_to_clear.append(i)

# Clear the old answer paragraphs
for i in paras_to_clear:
    doc.paragraphs[i].clear()

# Save
doc.save(r'C:\Users\theul\Downloads\06.03_Indigenous_Responses_FIXED.docx')
print("Saved to 06.03_Indigenous_Responses_FIXED.docx")
